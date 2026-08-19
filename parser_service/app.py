"""Lightweight CV text-extraction service for RecruitX.

POST /parse  -> multipart field `file` (pdf/docx), returns best-effort JSON.
GET  /health -> {status: ok}
"""

import io
import re
import unicodedata

import fitz
import docx
from fastapi import FastAPI, File, HTTPException, UploadFile

app = FastAPI(title="RecruitX CV Parser")


def norm(s: str) -> str:
    """Lowercase + strip accents, so French headers match English matcher rules."""
    s = unicodedata.normalize("NFD", s)
    return "".join(c for c in s if not unicodedata.combining(c)).lower()

SKILLS = [
    "Laravel", "PHP", "React", "TypeScript", "Vue", "Angular", "Java", "Spring Boot",
    "MySQL", "PostgreSQL", "Redis", "Docker", "Git", "AWS", "Python", "Node.js",
    "Tailwind CSS", "Figma", "Go", "C#", "JavaScript", "HTML", "CSS", "SQL",
    "jQuery", "Express", "MongoDB", "Next.js", "REST API", "GraphQL", "Django",
    "Flask", "FastAPI", "Machine Learning", "TensorFlow", "Kubernetes", "Terraform",
    "Jenkins", "CI/CD", "Linux", "Agile", "Scrum", "Redux", "Vite", "Webpack",
    "Selenium", "Jest", "Pandas", "NumPy", "C++", "Rust", "Ruby", "Kotlin", "Swift",
    "Flutter", "Azure", "GCP", "Microservices", "Docker Compose", "Nginx", "Elasticsearch",
]
SKILLS.sort(key=len, reverse=True)

SECTION_HEADERS = {
    "experience": [
        "work experience", "professional experience", "academic experience", "employment",
        "experience", "career history", "employment history",
        "expérience professionnelle", "expérience", "expériences académiques", "parcours professionnel", "emploi",
    ],
    "education": ["education", "qualifications", "formation", "études", "diplômes", "diplome", "éducation"],
    "skills": ["technical skills", "skills", "technologies", "tools", "compétences", "competences"],
}

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
# Generic international phone: optional +country code, digits grouped by 2-4
# with spaces/dots/dashes, at least 8 digits total (US, FR 0X XX XX XX XX,
# +33 6 XX XX XX XX, +1 555 123 4567, etc.). Matches any country, not just FR.
PHONE_RE = re.compile(
    r"(?:\+?\d{1,4}[ .\t-]?)?(?:\(?\d{1,4}\)?[ .\t-]?)?\d{1,4}[ .\t-]?\d{2,4}[ .\t-]?\d{2,4}(?:[ .\t-]?\d{2,4})?"
)
YEAR_RE = re.compile(r"(?:19|20)\d{2}")
DATE_RANGE_RE = re.compile(
    r"((?:[A-Za-z]+ )?(?:19|20)\d{2}(?:-\d{1,2})?)\s*(?:-|–|to)\s*((?:[A-Za-z]+ )?(?:19|20)\d{2}(?:-\d{1,2})?|present|current|now)",
    re.IGNORECASE,
)


def extract_text(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        doc = fitz.open(stream=data, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc)
        return text
    if name.endswith(".docx"):
        d = docx.Document(io.BytesIO(data))
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    raise ValueError("only pdf/docx are supported")


def sections(text: str) -> dict[str, str]:
    lines = text.splitlines()
    index = {}
    for i, line in enumerate(lines):
        ln = norm(line.strip())
        if not ln:
            continue
        for key, headers in SECTION_HEADERS.items():
            if key not in index and any(ln.startswith(norm(h)) for h in headers):
                index[key] = i
                break
    sorted_keys = sorted(index, key=index.get)
    result = {}
    for n, key in enumerate(sorted_keys):
        end = index[sorted_keys[n + 1]] if n + 1 < len(sorted_keys) else len(lines)
        result[key] = "\n".join(lines[index[key] + 1:end])
    return result


def extract_name(text: str, sec: dict) -> str | None:
    for line in text.splitlines()[:20]:
        line = line.strip()
        if not line or "@" in line or len(line) > 60:
            continue
        if line.lower() in {h for hs in SECTION_HEADERS.values() for h in hs}:
            continue
        words = line.split()
        if 2 <= len(words) <= 5 and all(w.isalpha() for w in words) and sum(w[:1].isupper() for w in words) >= 2:
            return line
    return None


def extract_skills(text: str) -> list[str]:
    found = []
    for skill in SKILLS:
        if re.search(r"(?<![\w])" + re.escape(skill) + r"(?![\w])", text, re.IGNORECASE) and skill not in found:
            found.append(skill)
    return found


def parse_year(value: str) -> str | None:
    m = YEAR_RE.search(value)
    return m.group(0) if m else None


def parse_date(value: str) -> str | None:
    m = re.search(r"(?:19|20)\d{2}(?:-\d{1,2})?", value)
    if not m:
        return None
    y = m.group(0)
    return y if len(y) == 4 else y


def extract_education(sec: str) -> list[dict]:
    result = []
    lines = [l.strip() for l in sec.splitlines() if l.strip()]
    for i, line in enumerate(lines):
        if not YEAR_RE.search(line):
            continue
        start_year = YEAR_RE.search(line).group(0)
        end_m = DATE_RANGE_RE.search(line)
        end_year = parse_date(end_m.group(2)) if end_m else None
        rest = line[end_m.end():].lstrip(" ,-–").strip() if end_m else line[YEAR_RE.search(line).end():].lstrip(" ,-–").strip()
        institution, degree = None, None
        if rest:
            parts = re.split(r"\s*(?:,|[-–]| at )\s*", rest)
            institution = parts[-1].strip()
            degree = ", ".join(p.strip() for p in parts[:-1]).strip() or None
        else:
            nxt = lines[i + 1].lstrip("•- \t").strip() if i + 1 < len(lines) else ""
            if nxt and not re.match(r"^[\w.\-']+,\s*[\w.\-']+$", nxt) and len(nxt) < 80:
                institution = nxt
            prev = lines[i - 1] if i > 0 else None
            if prev and not YEAR_RE.search(prev) and not re.match(r"^[\w.\-']+,\s*[\w.\-']+$", prev) \
                    and not prev.startswith("•") and len(prev) < 80 and not norm(prev).startswith("langu"):
                degree = prev
        if institution:
            result.append({
                "institution": institution,
                "degree": degree,
                "field_of_study": None,
                "start_date": start_year,
                "end_date": end_year,
            })
    return result[:8]


def extract_experience(sec: str) -> list[dict]:
    result = []
    lines = [l.strip() for l in sec.splitlines() if l.strip()]
    for i, line in enumerate(lines):
        m = DATE_RANGE_RE.search(line)
        if not m:
            continue
        start, end = m.group(1), m.group(2)
        rest = line[m.end():].lstrip(" ,-–").strip()
        position, company = None, None
        if rest:
            parts = re.split(r"\s*(?:,|-|@| at )\s*", rest)
            position = parts[0].strip() or None
            company = parts[1].strip() if len(parts) > 1 else None
        else:
            # French CVs put the job title ABOVE the date line, company below.
            prev = lines[i - 1] if i > 0 else None
            if prev and not norm(prev).startswith("mots-cl") and ":" not in prev and len(prev) < 80:
                position = prev
            company = lines[i + 1] if i + 1 < len(lines) else None
        result.append({
            "company_name": company,
            "position": position,
            "start_date": parse_date(start),
            "end_date": None if end.lower() in ("present", "current", "now") else parse_date(end),
            "is_current": end.lower() in ("present", "current", "now"),
        })
    return result[:10]


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/parse")
def parse(file: UploadFile = File(...)):
    data = file.file.read()
    try:
        text = extract_text(file.filename, data)
    except ValueError as e:
        raise HTTPException(status_code=422, detail=str(e))
    if not text.strip():
        raise HTTPException(status_code=422, detail="no extractable text")
    sec = sections(text)
    # Prefer the phone-like match with the most digits across all lines
    # (avoids grabbing zip codes or address numbers).
    best_phone, best_phone_digits = None, 0
    for line in text.splitlines():
        for m in PHONE_RE.finditer(line):
            digits = sum(c.isdigit() for c in m.group(0))
            if digits >= 8 and digits > best_phone_digits:
                best_phone, best_phone_digits = m.group(0).strip(), digits
    email = EMAIL_RE.search(text).group(0) if EMAIL_RE.search(text) else None
    return {
        "name": extract_name(text, sec),
        "email": email,
        "phone": best_phone,
        "address": None,
        "skills": extract_skills(text),
        "education": extract_education(sec.get("education", "")),
        "experiences": extract_experience(sec.get("experience", "")),
    }